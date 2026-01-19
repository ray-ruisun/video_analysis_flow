#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Word report generation module (Enhanced Version - SOTA 2026)

创建详细的 .docx 报告，包含:
- 执行摘要
- AI生成检测结果
- 跨视频共识分析 (含详细分布和置信度)
- 每个视频的详细分析 (含置信度)
- 数据表格
- 技术说明和建议

Models:
- CLIP (openai/clip-vit-large-patch14) - Scene Classification
- CLAP (laion/larger_clap_music_and_speech) - Audio Classification
- HuBERT (superb/hubert-large-superb-er) - Speech Emotion
- Whisper large-v3 (faster-whisper) - ASR
- YOLO26x (ultralytics) - Object Detection
- Deep-Fake-Detector-v2 - AI Detection
"""

import os
from pathlib import Path
from datetime import datetime
from typing import Dict, List, Any, Optional

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from utils import format_value


def set_cell_shading(cell, color: str):
    """设置单元格背景色"""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading)


def create_report_header(doc, title="Video Style Analysis Report", num_videos=1):
    """Add formatted header to document."""
    # Title
    title_para = doc.add_paragraph()
    title_run = title_para.add_run(title)
    title_run.bold = True
    title_run.font.size = Pt(24)
    title_run.font.color.rgb = RGBColor(31, 78, 120)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Subtitle with timestamp
    subtitle = doc.add_paragraph()
    timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    mode = "Multi-Video Comparison" if num_videos > 1 else "Single Video Analysis"
    subtitle_run = subtitle.add_run(f"{mode} ({num_videos} video{'s' if num_videos > 1 else ''}) | Generated: {timestamp}")
    subtitle_run.font.size = Pt(11)
    subtitle_run.font.color.rgb = RGBColor(100, 100, 100)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Technology stack notice
    tech = doc.add_paragraph()
    tech_run = tech.add_run(
        "🔧 SOTA 2026: CLIP | CLAP | HuBERT | Whisper large-v3 | YOLO26x | "
        "DeepFake-v2 | AIGC-Detector | Audio-Deepfake"
    )
    tech_run.font.size = Pt(9)
    tech_run.font.color.rgb = RGBColor(80, 120, 80)
    tech.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()


def add_section_header(doc, text, level=1):
    """Add formatted section header."""
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.bold = True
    
    if level == 1:
        run.font.size = Pt(16)
        run.font.color.rgb = RGBColor(31, 78, 120)
    elif level == 2:
        run.font.size = Pt(13)
        run.font.color.rgb = RGBColor(60, 100, 140)
    else:
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(80, 80, 80)
    
    return para


def add_executive_summary(doc, video_metrics, consensus):
    """Add executive summary section."""
    add_section_header(doc, "📋 Executive Summary", level=1)
    
    # Overview stats
    num_videos = len(video_metrics)
    total_duration = sum(
        m.get("visual", {}).get("duration", 0) for m in video_metrics
    )
    total_cuts = sum(
        m.get("visual", {}).get("cuts", 0) for m in video_metrics
    )
    
    # Check AI detection results
    ai_verdicts = [m.get("ai_detection", {}).get("verdict", "N/A") for m in video_metrics if m.get("ai_detection")]
    ai_summary = ""
    if ai_verdicts:
        real_count = sum(1 for v in ai_verdicts if v == "Real")
        ai_summary = f" AI Detection: {real_count}/{len(ai_verdicts)} classified as Real."
    
    doc.add_paragraph(
        f"This report analyzes {num_videos} video(s) with a combined duration of "
        f"{format_value(total_duration, '.1f')} seconds and {total_cuts} detected cuts.{ai_summary}"
    )
    
    # Key findings
    doc.add_paragraph()
    findings_para = doc.add_paragraph()
    findings_para.add_run("🔍 Key Findings:").bold = True
    
    findings = [
        f"Dominant camera angle: {consensus.get('camera_angle', 'N/A')}",
        f"Primary color tone: {consensus.get('hue_family', 'N/A')} with {consensus.get('saturation', 'N/A')} saturation",
        f"BGM style preference: {consensus.get('bgm_style', 'N/A')} ({consensus.get('bgm_mood', 'N/A')})",
        f"Editing pace: {format_value(consensus.get('cuts_per_minute'), '.1f')} cuts/minute",
        f"Scene type: {consensus.get('scene_category', 'N/A')}",
    ]
    
    # Add AI detection summary if available
    if ai_verdicts:
        for i, metrics in enumerate(video_metrics, 1):
            ai = metrics.get("ai_detection", {})
            if ai:
                verdict = ai.get('verdict', 'Unknown')
                confidence = ai.get('confidence', 0)
                findings.append(f"Video {i} AI verdict: {verdict} ({confidence:.1%})")
    
    for finding in findings:
        doc.add_paragraph(f"  • {finding}")
    
    doc.add_paragraph()


def add_video_source_table(doc, video_metrics):
    """Add source video information table."""
    add_section_header(doc, "📹 Source Videos", level=1)
    
    # Create table
    table = doc.add_table(rows=1, cols=5)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Header row
    headers = ['#', 'Video Name', 'Duration', 'FPS', 'Resolution']
    header_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        header_cells[i].text = header
        header_cells[i].paragraphs[0].runs[0].bold = True
        set_cell_shading(header_cells[i], 'E8F0FE')
    
    # Data rows
    for i, metrics in enumerate(video_metrics, 1):
        row = table.add_row().cells
        video_name = Path(metrics.get("path", f"Video {i}")).name
        duration = metrics.get("visual", {}).get("duration", 0)
        fps = metrics.get("visual", {}).get("fps", 0)
        total_frames = metrics.get("visual", {}).get("total_frames", 0)
        
        row[0].text = str(i)
        row[1].text = video_name[:30] + "..." if len(video_name) > 30 else video_name
        row[2].text = f"{format_value(duration, '.1f')}s"
        row[3].text = f"{format_value(fps, '.1f')}"
        row[4].text = f"{total_frames} frames"
    
    doc.add_paragraph()


def add_distribution_text(doc, detail: Dict, prefix: str = "    "):
    """Add distribution details as text."""
    distribution = detail.get('distribution', [])
    if distribution:
        for item in distribution[:5]:  # Top 5
            value = item.get('value', 'Unknown')
            count = item.get('count', 0)
            pct = item.get('percentage', 0)
            doc.add_paragraph(f"{prefix}• {value}: {count}次 ({pct}%)")


def add_consensus_section(doc, consensus):
    """Add enhanced cross-video consensus section."""
    add_section_header(doc, "🎯 I. Cross-Video Consensus Analysis", level=1)
    doc.add_paragraph(
        "Analysis of common patterns and majority preferences across all videos. "
        "Each metric shows the dominant value and detailed distribution."
    )
    
    # ========== Camera & Composition ==========
    add_section_header(doc, "📷 Camera & Composition", level=2)
    
    # Camera angle
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run("Camera Angle: ").bold = True
    p.add_run(f"{consensus.get('camera_angle', 'N/A')}")
    
    if consensus.get('camera_angle_detail'):
        add_distribution_text(doc, consensus['camera_angle_detail'])
    
    # Focal length
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run("Focal Length Tendency: ").bold = True
    p.add_run(f"{consensus.get('focal_length_tendency', 'N/A')}")
    
    if consensus.get('focal_length_detail'):
        add_distribution_text(doc, consensus['focal_length_detail'])
    
    # Camera motion
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run("Camera Motion: ").bold = True
    p.add_run(f"{consensus.get('camera_motion', 'N/A')}")
    
    if consensus.get('camera_motion_detail'):
        add_distribution_text(doc, consensus['camera_motion_detail'])
    
    # Composition
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run("Rule of Thirds: ").bold = True
    p.add_run(f"{consensus.get('composition_rule_of_thirds', 'N/A')}")
    
    if consensus.get('composition_detail'):
        add_distribution_text(doc, consensus['composition_detail'])
    
    # ========== Color & Lighting ==========
    add_section_header(doc, "🎨 Color & Lighting", level=2)
    
    # Hue
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run("Hue Family: ").bold = True
    p.add_run(f"{consensus.get('hue_family', 'N/A')}")
    
    if consensus.get('hue_detail'):
        add_distribution_text(doc, consensus['hue_detail'])
    
    # Saturation
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run("Saturation: ").bold = True
    p.add_run(f"{consensus.get('saturation', 'N/A')}")
    
    if consensus.get('saturation_detail'):
        add_distribution_text(doc, consensus['saturation_detail'])
    
    # Brightness
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run("Brightness: ").bold = True
    p.add_run(f"{consensus.get('brightness', 'N/A')}")
    
    if consensus.get('brightness_detail'):
        add_distribution_text(doc, consensus['brightness_detail'])
    
    # Contrast
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run("Contrast: ").bold = True
    p.add_run(f"{consensus.get('contrast', 'N/A')}")
    
    if consensus.get('contrast_detail'):
        add_distribution_text(doc, consensus['contrast_detail'])
    
    # CCT
    cct = consensus.get('cct')
    if cct:
        doc.add_paragraph()
        p = doc.add_paragraph()
        p.add_run("Color Temperature (CCT): ").bold = True
        p.add_run(f"≈ {format_value(cct, '.0f')} K")
    
    # Natural/Artificial light
    natural = consensus.get('natural_light_ratio')
    artificial = consensus.get('artificial_light_ratio')
    if natural is not None or artificial is not None:
        doc.add_paragraph()
        p = doc.add_paragraph()
        p.add_run("Lighting Balance: ").bold = True
        p.add_run(
            f"Natural {format_value(natural, '.1%')} / "
            f"Artificial {format_value(artificial, '.1%')}"
        )
    
    # ========== Pacing & Editing ==========
    add_section_header(doc, "✂️ Pacing & Editing", level=2)
    
    # Cuts per minute
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run("Cuts per Minute: ").bold = True
    p.add_run(f"{format_value(consensus.get('cuts_per_minute'), '.2f')}")
    
    # Average shot length
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run("Average Shot Length: ").bold = True
    p.add_run(f"{format_value(consensus.get('avg_shot_length'), '.2f')} seconds")
    
    # Transition type
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run("Transition Type: ").bold = True
    p.add_run(f"{consensus.get('transition_type', 'N/A')}")
    
    if consensus.get('transition_detail'):
        add_distribution_text(doc, consensus['transition_detail'])
    
    # ========== Audio & BGM ==========
    add_section_header(doc, "🎵 Audio & BGM (CLAP Analysis)", level=2)
    
    # BGM Style
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run("BGM Style: ").bold = True
    p.add_run(f"{consensus.get('bgm_style', 'N/A')}")
    
    if consensus.get('bgm_style_detail'):
        add_distribution_text(doc, consensus['bgm_style_detail'])
    
    # BGM Mood
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run("BGM Mood: ").bold = True
    p.add_run(f"{consensus.get('bgm_mood', 'N/A')}")
    
    if consensus.get('bgm_mood_detail'):
        add_distribution_text(doc, consensus['bgm_mood_detail'])
    
    # Tempo
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run("Tempo: ").bold = True
    p.add_run(f"≈ {format_value(consensus.get('tempo_bpm'), '.1f')} BPM")
    
    # Instruments
    instruments = consensus.get('bgm_instruments', [])
    if instruments:
        doc.add_paragraph()
        p = doc.add_paragraph()
        p.add_run("Detected Instruments: ").bold = True
        p.add_run(", ".join(instruments[:5]))
    
    # ========== Scene & Environment ==========
    add_section_header(doc, "🏠 Scene & Environment", level=2)
    
    # Scene category (CLIP)
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run("Scene Category (CLIP): ").bold = True
    p.add_run(f"{consensus.get('scene_category', 'N/A')}")
    
    if consensus.get('scene_category_detail'):
        add_distribution_text(doc, consensus['scene_category_detail'])
    
    # Countertop
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run("Countertop: ").bold = True
    p.add_run(
        f"{consensus.get('countertop_color', 'N/A')} / "
        f"{consensus.get('countertop_texture', 'N/A')}"
    )
    
    # YOLO Environment
    if consensus.get('yolo_available'):
        doc.add_paragraph()
        p = doc.add_paragraph()
        p.add_run("Environment (YOLO26): ").bold = True
        p.add_run(
            f"{consensus.get('yolo_environment', 'N/A')} - "
            f"{consensus.get('yolo_style', 'N/A')}"
        )
    
    doc.add_paragraph()


def add_per_video_section(doc, video_metrics, show_screenshots=True):
    """Add enhanced detailed per-video metrics section."""
    add_section_header(doc, "📊 II. Per-Video Detailed Analysis", level=1)
    
    for i, metrics in enumerate(video_metrics, 1):
        video_name = Path(metrics.get("path", f"Video {i}")).name
        
        # Video header
        doc.add_paragraph()
        video_header = doc.add_paragraph()
        run = video_header.add_run(f"Video {i}: {video_name}")
        run.bold = True
        run.font.size = Pt(13)
        run.font.color.rgb = RGBColor(60, 100, 140)
        
        # ========== Visual Analysis ==========
        vis = metrics.get("visual", {})
        if vis:
            add_section_header(doc, "Visual Analysis", level=3)
            
            # Create metrics table
            table = doc.add_table(rows=1, cols=2)
            table.style = 'Table Grid'
            
            visual_data = [
                ("Duration", f"{format_value(vis.get('duration'), '.2f')}s"),
                ("FPS", f"{format_value(vis.get('fps'), '.1f')}"),
                ("Sampled Frames", f"{vis.get('sampled_frames', 'N/A')}"),
                ("Camera Angle", vis.get('camera_angle', 'N/A')),
                ("Focal Length", vis.get('focal_length_tendency', 'N/A')),
                ("Hue Family", vis.get('hue_family', 'N/A')),
                ("Saturation", vis.get('saturation_band', 'N/A')),
                ("Brightness", vis.get('brightness_band', 'N/A')),
                ("Contrast", vis.get('contrast', 'N/A')),
                ("CCT", f"{format_value(vis.get('cct_mean'), '.0f')} K"),
                ("Cuts", f"{vis.get('cuts', 0)}"),
                ("Avg Shot Length", f"{format_value(vis.get('avg_shot_length'), '.2f')}s"),
                ("Transition Type", vis.get('transition_type', 'N/A')),
                ("Countertop", f"{vis.get('countertop_color', 'N/A')} / {vis.get('countertop_texture', 'N/A')}"),
            ]
            
            # Add header
            header_cells = table.rows[0].cells
            header_cells[0].text = "Metric"
            header_cells[1].text = "Value"
            for cell in header_cells:
                cell.paragraphs[0].runs[0].bold = True
                set_cell_shading(cell, 'F0F5FF')
            
            # Add data
            for metric, value in visual_data:
                row = table.add_row().cells
                row[0].text = metric
                row[1].text = str(value)
            
            # Scene categories
            scenes = vis.get('scene_categories', [])
            if scenes:
                doc.add_paragraph()
                p = doc.add_paragraph()
                p.add_run("Scene Classification (CLIP): ").bold = True
                for scene in scenes[:3]:
                    label = scene.get('label', 'Unknown')
                    prob = scene.get('probability', 0)
                    doc.add_paragraph(f"    • {label}: {prob:.1%}")
            
            # Camera angle distribution
            if vis.get('camera_angle_detail'):
                doc.add_paragraph()
                p = doc.add_paragraph()
                p.add_run("Camera Angle Distribution: ").bold = True
                add_distribution_text(doc, vis['camera_angle_detail'])
            
            # Hue distribution
            if vis.get('hue_detail'):
                doc.add_paragraph()
                p = doc.add_paragraph()
                p.add_run("Hue Distribution: ").bold = True
                add_distribution_text(doc, vis['hue_detail'])
        
        # ========== Audio Analysis ==========
        aud = metrics.get("audio", {})
        if aud and (aud.get("available") or aud.get("tempo_bpm")):
            add_section_header(doc, "Audio Analysis (CLAP)", level=3)
            
            table = doc.add_table(rows=1, cols=2)
            table.style = 'Table Grid'
            
            audio_data = [
                ("Tempo", f"{format_value(aud.get('tempo_bpm'), '.1f')} BPM"),
                ("Beats", f"{aud.get('num_beats', 0)}"),
                ("Percussive Ratio", f"{format_value(aud.get('percussive_ratio'), '.2f')}"),
                ("BGM Style", f"{aud.get('bgm_style', 'N/A')} ({format_value(aud.get('bgm_style_confidence'), '.1%')})"),
                ("Mood", f"{aud.get('mood', 'N/A')} ({format_value(aud.get('mood_confidence'), '.1%')})"),
                ("Key Signature", aud.get('key_signature', 'N/A')),
                ("Speech Ratio", f"{format_value(aud.get('speech_ratio'), '.2f')}"),
            ]
            
            header_cells = table.rows[0].cells
            header_cells[0].text = "Metric"
            header_cells[1].text = "Value"
            for cell in header_cells:
                cell.paragraphs[0].runs[0].bold = True
                set_cell_shading(cell, 'FFF5F0')
            
            for metric, value in audio_data:
                row = table.add_row().cells
                row[0].text = metric
                row[1].text = str(value)
            
            # BGM style detail
            if aud.get('bgm_style_detail') and aud['bgm_style_detail'].get('top_3'):
                doc.add_paragraph()
                p = doc.add_paragraph()
                p.add_run("BGM Style Scores (CLAP): ").bold = True
                for item in aud['bgm_style_detail']['top_3'][:5]:
                    if isinstance(item, (list, tuple)) and len(item) >= 2:
                        doc.add_paragraph(f"    • {item[0]}: {item[1]:.1%}")
            
            # Mood detail
            if aud.get('mood_detail') and aud['mood_detail'].get('top_3'):
                doc.add_paragraph()
                p = doc.add_paragraph()
                p.add_run("Mood Scores (CLAP): ").bold = True
                for item in aud['mood_detail']['top_3'][:5]:
                    if isinstance(item, (list, tuple)) and len(item) >= 2:
                        doc.add_paragraph(f"    • {item[0]}: {item[1]:.1%}")
            
            # Instruments
            instruments = aud.get('instruments', {})
            if instruments.get('detected_instruments'):
                doc.add_paragraph()
                p = doc.add_paragraph()
                p.add_run("Detected Instruments: ").bold = True
                p.add_run(", ".join(instruments['detected_instruments']))
        
        # ========== ASR Analysis ==========
        asr = metrics.get("asr", {})
        if asr and (asr.get("available") or asr.get("text")):
            add_section_header(doc, "Speech Analysis (Whisper + HuBERT)", level=3)
            
            table = doc.add_table(rows=1, cols=2)
            table.style = 'Table Grid'
            
            asr_data = [
                ("Word Count", f"{asr.get('num_words', 0)}"),
                ("Speech Rate", f"{format_value(asr.get('words_per_second'), '.2f')} w/s ({format_value(asr.get('words_per_minute'), '.1f')} wpm)"),
                ("Pace", asr.get('pace', 'N/A')),
                ("Pauses", f"{asr.get('num_pauses', 0)}"),
                ("Pause Style", asr.get('pause_style', 'N/A')),
                ("ASR Model", asr.get('implementation', 'Whisper')),
            ]
            
            header_cells = table.rows[0].cells
            header_cells[0].text = "Metric"
            header_cells[1].text = "Value"
            for cell in header_cells:
                cell.paragraphs[0].runs[0].bold = True
                set_cell_shading(cell, 'F0FFF0')
            
            for metric, value in asr_data:
                row = table.add_row().cells
                row[0].text = metric
                row[1].text = str(value)
            
            # Catchphrases
            catchphrases = asr.get('catchphrases', [])
            if catchphrases:
                doc.add_paragraph()
                p = doc.add_paragraph()
                p.add_run("Detected Catchphrases: ").bold = True
                for phrase in catchphrases[:10]:
                    doc.add_paragraph(f"    • \"{phrase}\"")
            
            # Prosody
            prosody = asr.get('prosody', {})
            if prosody:
                doc.add_paragraph()
                p = doc.add_paragraph()
                p.add_run("Prosody Analysis: ").bold = True
                doc.add_paragraph(f"    • Mean Pitch: {format_value(prosody.get('mean_pitch_hz'), '.1f')} Hz")
                doc.add_paragraph(f"    • Pitch Variation: {format_value(prosody.get('pitch_std'), '.1f')}")
                doc.add_paragraph(f"    • Tone: {prosody.get('tone', 'N/A')}")
                doc.add_paragraph(f"    • Style: {prosody.get('prosody_style', 'N/A')}")
            
            # Emotion
            emotion = asr.get('emotion', {})
            if emotion:
                doc.add_paragraph()
                p = doc.add_paragraph()
                p.add_run("Speech Emotion (HuBERT): ").bold = True
                doc.add_paragraph(f"    • Dominant: {emotion.get('dominant_emotion', 'N/A')} ({format_value(emotion.get('confidence'), '.1%')})")
                
                emotion_scores = emotion.get('emotion_scores', {})
                if emotion_scores:
                    doc.add_paragraph(f"    • All emotions:")
                    for emo, score in list(emotion_scores.items())[:5]:
                        doc.add_paragraph(f"        - {emo}: {score:.1%}")
            
            # Transcription preview
            text = asr.get('text', '')
            if text:
                doc.add_paragraph()
                p = doc.add_paragraph()
                p.add_run("Transcription (first 300 chars): ").bold = True
                preview = text[:300] + "..." if len(text) > 300 else text
                doc.add_paragraph(f"    \"{preview}\"")
        
        # ========== YOLO Analysis ==========
        yolo = metrics.get("yolo", {})
        if yolo and (yolo.get("available") or yolo.get("detection")):
            add_section_header(doc, "Object Detection (YOLO26x)", level=3)
            
            detection = yolo.get("detection", {})
            environment = yolo.get("environment", {})
            avg_conf = detection.get('avg_confidence', {})
            overall_conf = sum(avg_conf.values()) / max(len(avg_conf), 1) if avg_conf else 0
            
            table = doc.add_table(rows=1, cols=2)
            table.style = 'Table Grid'
            
            yolo_data = [
                ("Environment Type", f"{environment.get('environment_type', 'N/A')} ({environment.get('confidence', 0):.0%})"),
                ("Cooking Style", environment.get('cooking_style', 'N/A')),
                ("Appliance Tier", environment.get('appliance_tier', 'N/A')),
                ("Unique Objects", f"{detection.get('unique_objects', 0)}"),
                ("Total Detections", f"{detection.get('total_detections', 0)}"),
                ("Frames Processed", f"{detection.get('frames_processed', 0)}"),
                ("Avg Confidence", f"{overall_conf:.1%}"),
            ]
            
            header_cells = table.rows[0].cells
            header_cells[0].text = "Metric"
            header_cells[1].text = "Value"
            for cell in header_cells:
                cell.paragraphs[0].runs[0].bold = True
                set_cell_shading(cell, 'FFF0F5')
            
            for metric, value in yolo_data:
                row = table.add_row().cells
                row[0].text = metric
                row[1].text = str(value)
            
            # Detected objects with confidence
            object_counts = detection.get('object_counts', {})
            if object_counts:
                doc.add_paragraph()
                p = doc.add_paragraph()
                p.add_run("Detected Objects (with confidence): ").bold = True
                for obj, count in sorted(object_counts.items(), key=lambda x: x[1], reverse=True)[:10]:
                    conf = avg_conf.get(obj, 0)
                    doc.add_paragraph(f"    • {obj}: {count}x ({conf:.1%})")
            
            # Color analysis
            colors = yolo.get('colors', {})
            if colors:
                doc.add_paragraph()
                p = doc.add_paragraph()
                p.add_run("Object Colors: ").bold = True
                if colors.get('detailed_analysis'):
                    for obj, analysis in list(colors['detailed_analysis'].items())[:5]:
                        dominant = analysis.get('dominant', 'Unknown')
                        doc.add_paragraph(f"    • {obj}: {dominant}")
                        for item in analysis.get('distribution', [])[:3]:
                            if isinstance(item, dict):
                                doc.add_paragraph(f"        - {item.get('color', '?')}: {item.get('percentage', 0):.0%}")
                elif colors.get('dominant_colors'):
                    dom = colors.get('dominant_colors', [])
                    if isinstance(dom, list):
                        doc.add_paragraph(f"    • Dominant: {', '.join(str(c) for c in dom[:5])}")
                    elif isinstance(dom, dict):
                        for c, v in list(dom.items())[:5]:
                            doc.add_paragraph(f"    • {c}: {v}")
            
            # Material analysis
            materials = yolo.get('materials', {})
            if materials:
                doc.add_paragraph()
                p = doc.add_paragraph()
                p.add_run("Object Materials: ").bold = True
                if materials.get('detailed_analysis'):
                    for obj, analysis in list(materials['detailed_analysis'].items())[:5]:
                        dominant = analysis.get('dominant', 'Unknown')
                        doc.add_paragraph(f"    • {obj}: {dominant}")
                elif materials.get('dominant_materials'):
                    dom = materials.get('dominant_materials', [])
                    if isinstance(dom, list):
                        doc.add_paragraph(f"    • Dominant: {', '.join(str(m) for m in dom[:5])}")
                    elif isinstance(dom, dict):
                        for m, v in list(dom.items())[:5]:
                            doc.add_paragraph(f"    • {m}: {v}")
        
        # Add screenshot if available
        if show_screenshots:
            contact_sheet = vis.get("contact_sheet")
            if contact_sheet and os.path.exists(contact_sheet):
                try:
                    doc.add_paragraph()
                    p = doc.add_paragraph()
                    p.add_run("Contact Sheet:").bold = True
                    doc.add_picture(contact_sheet, width=Inches(6.5))
                    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
                except Exception as e:
                    doc.add_paragraph(f"    (Failed to insert contact sheet: {e})")
        
        doc.add_paragraph()  # Spacing between videos


def add_ai_detection_section(doc, video_metrics):
    """Add AI detection results section."""
    # Check if any video has AI detection results
    has_ai = any(m.get("ai_detection") for m in video_metrics)
    if not has_ai:
        return
    
    add_section_header(doc, "🤖 III. AI Generation Detection", level=1)
    doc.add_paragraph(
        "Multi-model ensemble detection for deepfakes, AI-generated content, and synthetic media. "
        "Uses weighted voting from multiple detection models for robust results."
    )
    
    for i, metrics in enumerate(video_metrics, 1):
        ai = metrics.get("ai_detection", {})
        if not ai:
            continue
        
        video_name = Path(metrics.get("path", f"Video {i}")).name
        
        # Video header
        doc.add_paragraph()
        video_header = doc.add_paragraph()
        run = video_header.add_run(f"Video {i}: {video_name}")
        run.bold = True
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(140, 60, 60)
        
        # Verdict and confidence
        verdict = ai.get('verdict', 'Unknown')
        confidence = ai.get('confidence', 0)
        
        doc.add_paragraph()
        p = doc.add_paragraph()
        p.add_run(f"Verdict: ").bold = True
        verdict_run = p.add_run(f"{verdict} ({confidence:.1%} confidence)")
        if verdict == "Real":
            verdict_run.font.color.rgb = RGBColor(0, 128, 0)
        elif verdict in ["Deepfake", "AI-Generated", "Synthetic"]:
            verdict_run.font.color.rgb = RGBColor(200, 0, 0)
        else:
            verdict_run.font.color.rgb = RGBColor(200, 150, 0)
        
        # Model scores table
        table = doc.add_table(rows=1, cols=4)
        table.style = 'Table Grid'
        
        headers = ['Model', 'Weight', 'Score', 'Status']
        header_cells = table.rows[0].cells
        for j, header in enumerate(headers):
            header_cells[j].text = header
            header_cells[j].paragraphs[0].runs[0].bold = True
            set_cell_shading(header_cells[j], 'FFE0E0')
        
        # Get weights from analysis_details
        weights = ai.get('analysis_details', {}).get('weights', {})
        
        models = [
            ("DeepFake-v2", weights.get('deepfake', 0.30), ai.get('deepfake_score', 0), ai.get('deepfake_available', False)),
            ("CLIP Synthetic", weights.get('clip', 0.20), ai.get('clip_synthetic_score', 0), ai.get('clip_available', False)),
            ("CLIP-Temporal", weights.get('temporal', 0.15), ai.get('temporal_score', 0), ai.get('temporal_available', False)),
            ("AIGC Detector", weights.get('aigc', 0.20), ai.get('aigc_score', 0), ai.get('aigc_available', False)),
            ("Audio Deepfake", weights.get('audio_deepfake', 0.10), ai.get('audio_deepfake_score', 0), ai.get('audio_deepfake_available', False)),
            ("Face Analysis", weights.get('face', 0.05), ai.get('no_face_ratio', 0), ai.get('face_available', False)),
        ]
        
        for model_name, weight, score, available in models:
            row = table.add_row().cells
            row[0].text = model_name
            row[1].text = f"{weight:.0%}"
            row[2].text = f"{score:.1%}"
            row[3].text = "✅" if available else "❌"
        
        # Additional details
        doc.add_paragraph()
        p = doc.add_paragraph()
        p.add_run("Detection Details: ").bold = True
        doc.add_paragraph(f"    • Faces Detected: {ai.get('faces_detected', 0)}")
        doc.add_paragraph(f"    • Frames Analyzed: {ai.get('frames_analyzed', 0)}")
        doc.add_paragraph(f"    • Frames with Faces: {ai.get('frames_with_faces', 0)}")
        doc.add_paragraph(f"    • Temporal Anomalies: {ai.get('temporal_anomalies', 0)}")
    
    doc.add_paragraph()


def add_technical_notes(doc):
    """Add technical notes section."""
    add_section_header(doc, "🔬 IV. Technical Notes", level=1)
    
    notes = [
        ("CLIP Scene Classification", 
         "Uses OpenAI CLIP ViT-L/14 for zero-shot scene classification. "
         "Supports custom scene descriptions for domain-specific classification."),
        
        ("CLAP Audio Analysis", 
         "Uses LAION CLAP for zero-shot audio classification. "
         "Classifies BGM style, mood, and instruments using audio-text contrastive learning."),
        
        ("HuBERT Speech Emotion", 
         "Uses HuBERT-large trained on SUPERB for speech emotion recognition. "
         "Classifies into angry, happy, neutral, sad."),
        
        ("Whisper large-v3", 
         "Full large-v3 Whisper model for highest accuracy ASR. Supports 100+ languages. "
         "Provides word-level timestamps for prosody analysis."),
        
        ("YOLO26x Object Detection", 
         "Latest YOLO26 extra-large model (Jan 2026) with highest accuracy. "
         "Uses COCO dataset classes for comprehensive object detection."),
        
        ("AI Detection Ensemble",
         "Multi-model approach: DeepFake-v2 (ViT, 92% acc), CLIP zero-shot synthetic detection, "
         "CLIP-based temporal consistency, AIGC detector, Audio deepfake detector. "
         "Final verdict from weighted ensemble voting."),
        
        ("Color Analysis",
         "HSV-based color classification with per-frame analysis. "
         "Returns distribution with confidence scores across sampled frames."),
        
        ("Shot Detection",
         "Uses PySceneDetect for content-aware shot boundary detection. "
         "Detects cuts, fades, and other transitions.")
    ]
    
    for title, description in notes:
        para = doc.add_paragraph()
        para.add_run(f"• {title}: ").bold = True
        para.add_run(description)
    
    doc.add_paragraph()


def add_references_section(doc):
    """Add references and citations."""
    add_section_header(doc, "📚 V. References & Models", level=1)
    
    references = [
        ("CLIP", "openai/clip-vit-large-patch14", "https://huggingface.co/openai/clip-vit-large-patch14"),
        ("CLAP", "laion/larger_clap_music_and_speech", "https://huggingface.co/laion/larger_clap_music_and_speech"),
        ("HuBERT", "superb/hubert-large-superb-er", "https://huggingface.co/superb/hubert-large-superb-er"),
        ("Whisper", "large-v3 via faster-whisper", "https://github.com/SYSTRAN/faster-whisper"),
        ("YOLO26", "yolo26x.pt via Ultralytics", "https://huggingface.co/collections/merve/yolo26-models"),
        ("DeepFake-v2", "prithivMLmods/Deep-Fake-Detector-v2-Model", "https://huggingface.co/prithivMLmods/Deep-Fake-Detector-v2-Model"),
        ("AIGC Detector", "umm-maybe/AI-image-detector", "https://huggingface.co/umm-maybe/AI-image-detector"),
        ("Audio Deepfake", "MelodyMachine/Deepfake-audio-detection", "https://huggingface.co/MelodyMachine/Deepfake-audio-detection"),
        ("PySceneDetect", "Shot Boundary Detection", "https://www.scenedetect.com/"),
        ("librosa", "Audio Feature Extraction", "https://librosa.org/"),
        ("OpenCV", "Computer Vision", "https://opencv.org/"),
    ]
    
    for name, model, url in references:
        doc.add_paragraph(f"• {name}: {model} - {url}")


def generate_word_report(video_metrics, consensus, output_path, 
                        show_screenshots=True, show_references=True):
    """
    Generate comprehensive Word document report.
    
    Args:
        video_metrics: List of per-video metric dictionaries
        consensus: Cross-video consensus dictionary
        output_path: Path to save .docx file
        show_screenshots: Whether to include contact sheets
        show_references: Whether to include references section
        
    Returns:
        str: Path to generated report
    """
    # Create document
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(10)
    
    num_videos = len(video_metrics)
    
    # Build report sections
    create_report_header(doc, num_videos=num_videos)
    add_executive_summary(doc, video_metrics, consensus)
    add_video_source_table(doc, video_metrics)
    add_consensus_section(doc, consensus)
    add_per_video_section(doc, video_metrics, show_screenshots)
    add_ai_detection_section(doc, video_metrics)  # New AI detection section
    add_technical_notes(doc)
    
    if show_references:
        add_references_section(doc)
    
    # Save document
    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    
    return str(output_path)
